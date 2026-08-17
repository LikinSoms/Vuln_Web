from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\cyber\Northstar Market - VAPT Report.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(23, 32, 42)
MUTED = RGBColor(90, 100, 112)
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "DADFE6"
PALE_BLUE = "E8EEF5"
PALE_RED = "FDEDEC"
PALE_ORANGE = "FFF4E6"
PALE_YELLOW = "FFF8E1"
PALE_GREEN = "EAF7EF"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def paragraph_border_bottom(paragraph, color="2E74B5", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="DADFE6", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)


def cell_text(cell, text, bold=False, color=INK, size=10.2, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.05)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    return paragraph


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_GRAY, header_color=INK):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        shade_cell(header_cells[i], header_fill)
        cell_text(header_cells[i], header, bold=True, color=header_color, size=9.8)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell_text(cells[i], value, size=9.6)
    after = doc.add_paragraph()
    set_paragraph_spacing(after, before=0, after=4)
    return table


def add_paragraph(doc, text="", bold=False, italic=False, color=INK, size=11, after=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_numbered(doc, items):
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=4, line_spacing=1.167)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.append(num_pr)
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_bullets(doc, items):
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "-")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=4, line_spacing=1.167)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.append(num_pr)
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_step_table(doc, steps):
    rows = [[f"Step {index}", step] for index, step in enumerate(steps, 1)]
    add_table(
        doc,
        ["Step", "Action"],
        rows,
        [1050, 8310],
        header_fill=LIGHT_GRAY,
    )


def add_screenshot_placeholder(doc, title, guidance):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="BFC7D1", size="8")
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F9FB")
    set_cell_margins(cell, top=180, bottom=180, start=120, end=120)
    row = table.rows[0]
    row.height = Inches(1.25)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=2, line_spacing=1.1)
    r = p.add_run(f"{title}\n")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(guidance)
    set_run_font(r, size=9.5, color=MUTED, italic=True)
    add_paragraph(doc, "", after=4)


def add_finding(doc, finding):
    add_heading(doc, finding["heading"], 2)
    summary_rows = [
        ("CWE", finding["cwe"]),
        ("Severity", finding["severity"]),
        ("Estimated CVSS", finding["cvss"]),
        ("Vulnerable Instance", finding["endpoint"]),
        ("Affected Component", finding["component"]),
    ]
    table = add_table(
        doc,
        ["Field", "Value"],
        summary_rows,
        [2200, 7160],
        header_fill=PALE_BLUE,
    )
    for row in table.rows[1:]:
        if row.cells[0].text == "Severity":
            shade_cell(row.cells[1], finding["severity_fill"])
            row.cells[1].paragraphs[0].runs[0].bold = True

    add_heading(doc, "Description", 3)
    add_paragraph(doc, finding["description"])

    add_heading(doc, "Business Impact", 3)
    add_bullets(doc, finding["impact"])

    add_heading(doc, "Proof of Concept Steps", 3)
    add_step_table(doc, finding["steps"])
    add_screenshot_placeholder(
        doc,
        f"Screenshot Placeholder - {finding['short_name']} Evidence",
        finding["screenshot_guidance"],
    )

    add_heading(doc, "Recommendation", 3)
    add_bullets(doc, finding["recommendations"])

    add_heading(doc, "Retest Criteria", 3)
    add_bullets(doc, finding["retest"])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.font.color.rgb = INK
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(hp, after=0)
    run = hp.add_run("Northstar Market VAPT Report | Local Sandbox")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(fp, after=0)
    run = fp.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    add_field(fp, "PAGE")


def build_doc():
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Northstar Market - VAPT Report"
    doc.core_properties.subject = "Vulnerability Assessment and Penetration Testing Report"
    doc.core_properties.author = "[Your Name]"

    # Cover page
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=0, after=18)
    run = title.add_run("Northstar Market - VAPT Report")
    set_run_font(run, size=26, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=28)
    run = subtitle.add_run("Vulnerability Assessment and Penetration Testing on a Local E-Commerce Sandbox")
    set_run_font(run, size=14, color=INK, italic=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(meta, after=8)
    run = meta.add_run("August 17, 2026")
    set_run_font(run, size=12, color=INK)

    byline = doc.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(byline, after=8)
    run = byline.add_run("Prepared by: [Your Name]")
    set_run_font(run, size=12, color=INK, italic=True)

    doc.add_page_break()

    add_heading(doc, "Table of Contents", 1)
    contents = [
        "1. Project Summary",
        "   1.1 Executive Summary",
        "   1.2 Scope",
        "   1.3 Project Background & Context",
        "2. Assessment Methodology",
        "   2.1 Tools Used",
        "3. Assessment Findings",
        "   3.1 Summary of Vulnerabilities",
        "   3.2 Detailed Findings",
        "4. Assessment Recommendations",
        "5. Appendix - Screenshot Checklist",
    ]
    for item in contents:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=2, line_spacing=1.10)
        if item.startswith("   "):
            p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(item)
        set_run_font(run, size=11.5 if not item.startswith("   ") else 10.8, color=INK)

    doc.add_page_break()

    add_heading(doc, "1. Project Summary", 1)
    add_heading(doc, "1.1 Executive Summary", 2)
    add_paragraph(
        doc,
        "This document summarizes the vulnerability assessment and penetration testing notes for Northstar Market, "
        "a local Express and SQLite e-commerce sandbox designed for controlled Burp Suite practice. The application "
        "provides product browsing, account registration, login, cart checkout, order history, and invoice retrieval."
    )
    add_paragraph(
        doc,
        "The assessment identified three main vulnerabilities in the target application: SQL Injection in the login "
        "workflow, Broken Object Level Authorization in invoice retrieval, and checkout price manipulation through "
        "trusted client-side price parameters. Screenshots are intentionally left as placeholders so evidence captured "
        "during manual Burp Suite testing can be inserted later."
    )

    add_heading(doc, "1.1.1 Project Objectives", 3)
    add_bullets(doc, [
        "Document the security posture of the local Northstar Market sandbox in a VAPT report format.",
        "Capture the three primary vulnerabilities with affected routes, impact, reproduction steps, and fixes.",
        "Provide a clean Word document structure where Burp Suite screenshots can be added after testing.",
        "Use the attached JPetStore Demo VAPT report as a structural reference only."
    ])

    add_heading(doc, "1.1.2 Summary of Findings", 3)
    add_table(
        doc,
        ["Sl. No.", "Vulnerability", "CWE", "Severity", "Affected Endpoint"],
        [
            ["1", "SQL Injection on Login", "CWE-89", "Critical", "POST /api/login"],
            ["2", "Broken Object Level Authorization / IDOR", "CWE-639", "High", "GET /api/orders/:id/invoice"],
            ["3", "Parameter Tampering / Price Manipulation", "CWE-494", "High", "POST /api/checkout"],
        ],
        [850, 3000, 1100, 1200, 3210],
        header_fill=PALE_BLUE,
    )

    add_heading(doc, "1.2 Scope", 2)
    add_table(
        doc,
        ["Scope", "Scope Type", "Target", "Assessment Date"],
        [
            ["Northstar Market", "Local Web Application", "http://localhost:3000", "August 17, 2026"],
            ["Source Folder", "Code Review Reference", r"D:\E-commerce", "August 17, 2026"],
        ],
        [2100, 2100, 2900, 2260],
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "1.3 Project Background & Context", 2)
    add_paragraph(
        doc,
        "Northstar Market is a minimal e-commerce training application built with Node.js, Express, SQLite, and vanilla "
        "HTML/CSS/JavaScript. It serves the frontend directly from Express static folders and stores users, products, "
        "orders, and invoice line items in a local SQLite database."
    )
    add_heading(doc, "1.3.1 Functions/Components", 3)
    add_bullets(doc, [
        "Product catalog with dynamic items and thumbnail images.",
        "Cart management using browser-side JavaScript and local storage.",
        "User login and registration pages.",
        "Authenticated checkout flow that creates local SQLite orders.",
        "Order history and invoice retrieval by order ID.",
    ])
    add_heading(doc, "1.3.2 Goals of the Assessment", 3)
    add_paragraph(
        doc,
        "The goal of this assessment is to practice identifying, exploiting, documenting, and remediating common web "
        "application vulnerabilities in a fully authorized local environment. Burp Suite can be used to intercept "
        "requests, modify parameters, and capture evidence for each vulnerability."
    )

    add_heading(doc, "2. Assessment Methodology", 1)
    add_paragraph(
        doc,
        "Testing focused on the user-facing application workflows and their corresponding HTTP requests. Authentication, "
        "invoice retrieval, and checkout were prioritized because they handle credentials, authorization decisions, and "
        "order totals."
    )
    add_heading(doc, "2.1 Tools Used", 2)
    add_bullets(doc, [
        "Burp Suite for proxy interception, request replay, and parameter tampering.",
        "Browser developer tools for client-side request and response inspection.",
        "Node.js and Express runtime for local application execution.",
        "SQLite database for local application state validation.",
        "Manual source review of server-side routes and data flow.",
    ])

    add_heading(doc, "3. Assessment Findings", 1)
    add_heading(doc, "3.1 Vulnerabilities", 2)
    add_paragraph(
        doc,
        "The following findings represent the three main intentionally exposed weaknesses in the sandbox. Each finding "
        "includes a vulnerable route, impact, proof-of-concept workflow, recommended fix, and retest criteria."
    )

    findings = [
        {
            "heading": "3.1.1 SQL Injection on Login",
            "short_name": "SQL Injection",
            "cwe": "CWE-89",
            "severity": "Critical",
            "severity_fill": PALE_RED,
            "cvss": "9.1",
            "endpoint": "POST /api/login",
            "component": "Authentication",
            "description": (
                "The login endpoint builds the SQL query by concatenating the submitted username into the query string. "
                "Because user-controlled input is interpreted as part of the SQL statement, an attacker can alter the "
                "query logic and bypass normal authentication checks."
            ),
            "impact": [
                "Unauthorized login as an existing user.",
                "Potential exposure of order history and invoice data after account takeover.",
                "Loss of confidence in authentication controls and auditability.",
            ],
            "steps": [
                "Open the application through Burp Suite and browse to /login.html.",
                "Submit any password while placing a SQL condition in the username field, such as ' OR 1=1 -- .",
                "Send the request through Burp Repeater and observe whether the response creates a valid session.",
                "Navigate back to the store and confirm the application reports an authenticated user.",
            ],
            "screenshot_guidance": (
                "Insert the intercepted login request, injected username parameter, and successful authenticated response."
            ),
            "recommendations": [
                "Use parameterized SQLite queries for username and password lookup.",
                "Return consistent authentication failure messages and log suspicious patterns server-side.",
                "Add rate limiting and account lockout controls to reduce automated login abuse.",
            ],
            "retest": [
                "The same payload should return an invalid credentials response.",
                "SQL metacharacters in username or password should be treated as literal input.",
                "Application logs should show the failed attempt without leaking SQL errors to the client.",
            ],
        },
        {
            "heading": "3.1.2 Broken Object Level Authorization / IDOR",
            "short_name": "IDOR",
            "cwe": "CWE-639",
            "severity": "High",
            "severity_fill": PALE_ORANGE,
            "cvss": "8.1",
            "endpoint": "GET /api/orders/:id/invoice",
            "component": "Order Invoice Retrieval",
            "description": (
                "The invoice endpoint accepts a sequential order ID and returns the invoice for that order without "
                "verifying that the authenticated session owns the order. A logged-in user can change the order ID in "
                "the URL and request invoices belonging to other users."
            ),
            "impact": [
                "Unauthorized disclosure of another user's order details.",
                "Exposure of customer identity fields shown in invoice responses.",
                "Enumeration risk because order identifiers are predictable integers.",
            ],
            "steps": [
                "Log in as alice and view an invoice from the Your Orders panel.",
                "Send the invoice request to Burp Repeater.",
                "Change the order ID path value to another sequential ID, such as /api/orders/2/invoice.",
                "Send the modified request and observe whether another customer's invoice is returned.",
            ],
            "screenshot_guidance": (
                "Insert the original invoice request, the modified order ID request, and the response showing another user's invoice."
            ),
            "recommendations": [
                "Filter invoice lookups by both order ID and the authenticated user's session ID.",
                "Return 404 or 403 when the order does not belong to the current user.",
                "Consider non-sequential public order references for user-facing invoice links.",
            ],
            "retest": [
                "A user should only retrieve invoices where orders.user_id matches the session user ID.",
                "Changing the order ID to another user's order should fail.",
                "Order history and invoice routes should enforce the same authorization policy.",
            ],
        },
        {
            "heading": "3.1.3 Parameter Tampering / Price Manipulation",
            "short_name": "Price Manipulation",
            "cwe": "CWE-494",
            "severity": "High",
            "severity_fill": PALE_ORANGE,
            "cvss": "8.2",
            "endpoint": "POST /api/checkout",
            "component": "Checkout and Order Pricing",
            "description": (
                "The checkout endpoint trusts the item price submitted in the frontend request body. Because the price "
                "is controlled by the client, a user can intercept the checkout request and replace item prices before "
                "the order is saved."
            ),
            "impact": [
                "Creation of orders with fraudulent totals.",
                "Revenue loss and corrupted order records.",
                "Failure of backend price integrity controls in the purchase workflow.",
            ],
            "steps": [
                "Log in and add a product, such as Bluetooth Speaker, to the cart.",
                "Click Checkout while Burp Suite is intercepting the POST /api/checkout request.",
                "Modify the JSON item price value to a lower number, such as 0.01, and forward the request.",
                "Observe that the order total reflects the tampered client-provided price.",
            ],
            "screenshot_guidance": (
                "Insert the intercepted checkout JSON before and after price modification, plus the resulting low-total order response."
            ),
            "recommendations": [
                "Ignore client-submitted item prices during checkout.",
                "Look up authoritative product prices from the database on the server before calculating totals.",
                "Store immutable order line item prices calculated by the backend at checkout time.",
            ],
            "retest": [
                "Changing price values in the request body should not affect the final order total.",
                "The server should calculate totals using current database prices.",
                "Order records should preserve server-calculated line item subtotals.",
            ],
        },
    ]

    for index, finding in enumerate(findings):
        if index == 1:
            doc.add_page_break()
        add_finding(doc, finding)

    add_heading(doc, "3.1.4 Severity Analysis of Identified Vulnerabilities", 2)
    add_table(
        doc,
        ["Sl. No.", "Vulnerability", "Estimated CVSS", "Severity"],
        [
            ["1", "SQL Injection on Login", "9.1", "Critical"],
            ["2", "Broken Object Level Authorization / IDOR", "8.1", "High"],
            ["3", "Parameter Tampering / Price Manipulation", "8.2", "High"],
        ],
        [900, 4900, 1700, 1860],
        header_fill=PALE_BLUE,
    )

    add_heading(doc, "4. Assessment Recommendations", 1)
    add_paragraph(
        doc,
        "The application should be remediated by enforcing server-side trust boundaries. Authentication, authorization, "
        "and pricing decisions must be made on the backend using validated session context and authoritative database "
        "records."
    )
    add_table(
        doc,
        ["Priority", "Action", "Owner", "Validation"],
        [
            ["P1", "Replace raw login SQL with parameterized SQLite queries.", "Backend", "SQL injection payload fails."],
            ["P1", "Require order ownership in invoice lookup queries.", "Backend", "Cross-user invoice IDs return 403 or 404."],
            ["P1", "Recalculate checkout totals from database prices.", "Backend", "Tampered price values are ignored."],
            ["P2", "Add automated regression tests for auth, invoice access, and checkout totals.", "Engineering", "Tests fail on vulnerable behavior and pass after patches."],
            ["P2", "Improve session hardening, logging, and request validation.", "Engineering", "Security events are logged and invalid input is handled consistently."],
        ],
        [1000, 3900, 1500, 2960],
        header_fill=PALE_GREEN,
    )

    add_heading(doc, "5. Appendix - Screenshot Checklist", 1)
    add_paragraph(
        doc,
        "Use this checklist when adding screenshots after Burp Suite testing. Keep screenshots close to the matching "
        "finding section and redact anything outside the local sandbox if it appears in the browser or proxy."
    )
    add_table(
        doc,
        ["Finding", "Screenshot to Add", "Suggested Placement"],
        [
            ["Application Overview", "Home page showing product catalog and cart.", "Project Background & Context"],
            ["SQL Injection", "Burp request with injected username and successful login response.", "Finding 3.1.1"],
            ["IDOR", "Invoice request where the order ID is changed to another user's order.", "Finding 3.1.2"],
            ["Price Manipulation", "Checkout JSON with modified price and resulting order total.", "Finding 3.1.3"],
            ["Retest", "Patched behavior showing failed exploit attempts.", "Assessment Recommendations or separate retest appendix"],
        ],
        [2100, 4600, 2660],
        header_fill=PALE_YELLOW,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
